"""TDD: the docx_v2 adapter must register every embedded image it links
to in the markdown as an ``AssetRef``.

Before this fix the adapter wrote ``![](assets/img_N.png)`` markdown links
and saved the image bytes, but returned ``assets=()``. The deterministic
core's H5 check ("markdown references not in AssetRefs") then failed for
*any* DOCX containing an embedded image, which is virtually every real
engineering document.
"""

from __future__ import annotations

import io
import re
import struct
import zlib
from pathlib import Path

import docx as _docx

from kb_extract.adapters.docx_v2 import DocxV2Adapter


def _png_bytes(width: int = 96, height: int = 96) -> bytes:
    """Build a minimal but valid (and >= 1 KiB) PNG in pure Python.

    Uses pseudo-random pixel data so the zlib-compressed payload stays
    comfortably above ``MIN_IMAGE_BYTES`` (1 KiB); a flat-colour image
    would compress below the threshold and be skipped by ``save_image``.
    """

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = bytearray()
    seed = 1234567
    for _ in range(height):
        raw.append(0)  # filter byte
        for _ in range(width * 3):
            seed = (seed * 1103515245 + 12345) & 0x7FFFFFFF
            raw.append((seed >> 16) & 0xFF)
    idat = zlib.compress(bytes(raw), 6)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", idat)
        + chunk(b"IEND", b"")
    )


def _make_docx_with_image(path: Path) -> None:
    doc = _docx.Document()
    doc.add_heading("Requirements", level=1)
    doc.add_paragraph("Some body text describing a requirement.")
    doc.add_picture(io.BytesIO(_png_bytes()))
    doc.save(str(path))


def test_docx_v2_registers_embedded_image_as_asset(tmp_path: Path) -> None:
    src = tmp_path / "with_image.docx"
    _make_docx_with_image(src)

    result = DocxV2Adapter().extract(src, tmp_path)

    # Every assets/... path referenced in the markdown must be backed by an
    # AssetRef -- this is exactly what the H5 check enforces.
    md_refs = set(re.findall(r"!\[\]\((assets/[^)]+)\)", result.markdown))
    assert md_refs, "fixture should produce at least one image markdown link"

    asset_paths = {a.rel_path for a in result.assets}
    assert md_refs <= asset_paths, (
        f"markdown references {md_refs} not registered as AssetRefs "
        f"{asset_paths}"
    )

    for asset in result.assets:
        assert asset.kind == "image"
        assert asset.sha256 and len(asset.sha256) == 64
