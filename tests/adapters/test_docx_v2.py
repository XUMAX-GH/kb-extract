"""Tests for DocxV2Adapter (v0.8.0 parser v2).

Focuses on the two highest-impact improvements over the legacy DOCX
adapter (spec sec.4.1):
  1. Merged cells via gridSpan / vMerge XML -> HTML <table>
  2. Embedded images via a:blip XML -> assets/ + ![](...) markdown

OMML formulas, list nesting, and run-level formatting are deferred to
follow-up commits and are not tested here.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from kb_extract.adapters.docx_v2 import DocxV2Adapter
from kb_extract.hardness import assert_invariants

pytestmark = pytest.mark.disable_socket


# --- fixtures ----------------------------------------------------------------


def _make_simple_docx(path: Path) -> Path:
    """Plain docx with a heading + paragraph + 2x2 table (no merges)."""
    from docx import Document
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "a"
    t.cell(0, 1).text = "b"
    t.cell(1, 0).text = "c"
    t.cell(1, 1).text = "d"
    doc.save(str(path))
    return path


def _make_merged_cells_docx(path: Path) -> Path:
    """2x3 grid where the top-left two cells are horizontally merged."""
    from docx import Document
    doc = Document()
    doc.add_heading("Merged Cells", level=1)
    t = doc.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "top"
    t.cell(0, 1).text = "x"  # will be merged into (0,0)
    t.cell(0, 2).text = "right"
    t.cell(1, 0).text = "a"
    t.cell(1, 1).text = "b"
    t.cell(1, 2).text = "c"
    # Horizontal merge: cell(0,0) absorbs cell(0,1)
    t.cell(0, 0).merge(t.cell(0, 1))
    doc.save(str(path))
    return path


def _make_vertical_merge_docx(path: Path) -> Path:
    """3x2 grid where the left column rows 0-1 are vertically merged."""
    from docx import Document
    doc = Document()
    doc.add_heading("Vertical Merge", level=1)
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "tall"
    t.cell(0, 1).text = "x1"
    t.cell(1, 0).text = "y"  # will be merged up
    t.cell(1, 1).text = "x2"
    t.cell(2, 0).text = "bottom"
    t.cell(2, 1).text = "x3"
    t.cell(0, 0).merge(t.cell(1, 0))
    doc.save(str(path))
    return path


def _make_embedded_image_docx(path: Path, tmp_path: Path) -> Path:
    """DOCX with a heading and one embedded PNG (≥1KB)."""
    import secrets

    from docx import Document
    from docx.shared import Inches
    from PIL import Image as PILImage

    # Generate a noise-PNG large enough to clear the 1 KiB threshold even
    # after PNG compression. 200x200 of pseudo-random bytes works.
    png_path = tmp_path / "_inline.png"
    noise = secrets.token_bytes(200 * 200 * 3)
    img = PILImage.frombytes("RGB", (200, 200), noise)
    img.save(png_path, format="PNG")
    assert png_path.stat().st_size >= 1024, f"fixture PNG only {png_path.stat().st_size}B"

    doc = Document()
    doc.add_heading("With Image", level=1)
    doc.add_paragraph("Before image.")
    doc.add_picture(str(png_path), width=Inches(1.0))
    doc.add_paragraph("After image.")
    doc.save(str(path))
    return path


# --- baseline parity with legacy --------------------------------------------


def test_v2_extracts_headings_and_body(tmp_path: Path) -> None:
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert "# Title" in res.markdown
    assert "Body text." in res.markdown


def test_v2_passes_hardness_invariants(tmp_path: Path) -> None:
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert_invariants(res, src, out_dir, total_pages=res.index.page_end)


def test_v2_deterministic(tmp_path: Path) -> None:
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    r1 = DocxV2Adapter().extract(src, out_dir)
    out_dir2 = tmp_path / "out2"
    (out_dir2 / "assets").mkdir(parents=True)
    r2 = DocxV2Adapter().extract(src, out_dir2)
    assert r1.markdown == r2.markdown


def test_v2_adapter_metadata(tmp_path: Path) -> None:
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert res.meta.adapter_name == "docx_v2"
    assert DocxV2Adapter.extensions == (".docx",)


# --- merged cells ------------------------------------------------------------


def test_v2_emits_html_table_not_markdown_pipe(tmp_path: Path) -> None:
    """v2 always emits raw HTML <table> for tables, so merged cells survive."""
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert "<table>" in res.markdown
    assert "</table>" in res.markdown


def test_v2_horizontal_merge_emits_colspan(tmp_path: Path) -> None:
    src = _make_merged_cells_docx(tmp_path / "h.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert 'colspan="2"' in res.markdown
    # Merged cell content includes the original "top" text
    assert "top" in res.markdown
    # The second-row cells should NOT be merged
    assert "<td>a</td>" in res.markdown
    assert "<td>b</td>" in res.markdown
    assert "<td>c</td>" in res.markdown


def test_v2_vertical_merge_emits_rowspan(tmp_path: Path) -> None:
    src = _make_vertical_merge_docx(tmp_path / "v.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert 'rowspan="2"' in res.markdown
    # The merged content "tall" appears exactly once (continuation cell dropped)
    # Note: python-docx concatenates merged text differently across raw-XML
    # vs. python-docx text path, so we check the "tall" substring count.
    assert res.markdown.count("tall") == 1
    assert "<td>bottom</td>" in res.markdown


def test_v2_table_anchor_present(tmp_path: Path) -> None:
    src = _make_simple_docx(tmp_path / "s.docx")
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    assert '<a id="tbl-0001"></a>' in res.markdown
    assert len(res.tables) == 1


# --- embedded images ---------------------------------------------------------


def test_v2_extracts_embedded_image_to_assets(tmp_path: Path) -> None:
    src = _make_embedded_image_docx(tmp_path / "img.docx", tmp_path)
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    # Markdown contains an image link
    assert "![](assets/" in res.markdown
    # File exists in out_dir/assets/
    assets = list((out_dir / "assets").glob("img_*.png"))
    assert len(assets) >= 1
    assert assets[0].stat().st_size >= 1024


def test_v2_image_appears_between_paragraphs(tmp_path: Path) -> None:
    src = _make_embedded_image_docx(tmp_path / "img.docx", tmp_path)
    out_dir = tmp_path / "out"
    (out_dir / "assets").mkdir(parents=True)
    res = DocxV2Adapter().extract(src, out_dir)
    md = res.markdown
    pos_before = md.find("Before image.")
    pos_img = md.find("![](assets/")
    pos_after = md.find("After image.")
    assert -1 < pos_before < pos_img < pos_after
