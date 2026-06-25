"""DOCX v2 adapter — MinerU-inspired (v0.8.0 parser v2, spec §4.1).

Scope of this commit:
- Merged cells via ``gridSpan`` (colspan) and ``vMerge`` (rowspan) XML →
  HTML ``<table>`` via :mod:`._table_utils`. Markdown pipe tables cannot
  represent merged cells, so we always emit raw HTML for tables.
- Embedded images via ``<a:blip r:embed="...">`` XML → ``assets/`` via
  :mod:`._image_utils`, with ``![](assets/...)`` markdown links inserted
  at the correct paragraph position (document order preserved).

Out of scope (deferred to follow-up commits per spec §4.1):
- OMML formula → LaTeX conversion (will emit ``[formula]`` placeholder
  until then; current commit just keeps any inline text content).
- List nesting (numPr levels) — currently falls back to plain paragraph
  rendering.
- Run-level formatting (bold/italic/links).

Headings, language detection, page count, anchor scheme, hardness
guarantees match the legacy ``DocxAdapter``.
"""
from __future__ import annotations

import hashlib
from pathlib import Path

import docx as _docx
import langdetect
from docx.oxml.ns import qn

from ..contracts import AssetRef, ExtractionResult, SectionNode, TableRef
from ._common import make_meta, sha256_bytes
from ._image_utils import save_image
from ._table_utils import CellInfo, cells_to_html

_DOCX_VERSION = getattr(_docx, "__version__", "unknown")

_HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
}

_QN_P = qn("w:p")
_QN_TBL = qn("w:tbl")
_QN_TR = qn("w:tr")
_QN_TC = qn("w:tc")
_QN_TC_PR = qn("w:tcPr")
_QN_GRID_SPAN = qn("w:gridSpan")
_QN_V_MERGE = qn("w:vMerge")
_QN_BLIP = qn("a:blip")
_QN_EMBED = qn("r:embed")


def _detect_lang(text: str) -> str:
    try:
        return langdetect.detect(text) if text.strip() else "und"
    except Exception:
        return "und"


def _build_cell_grid(table) -> list[list[CellInfo]]:
    """Convert a python-docx Table into a phantom-free grid for cells_to_html.

    Iterates the raw ``<w:tr>``/``<w:tc>`` XML rather than python-docx's
    ``row.cells`` (which synthesises virtual cells for merged regions and
    is unreliable). For each tc:

      - ``gridSpan`` → colspan
      - ``vMerge`` with ``val="restart"`` → start of a vertical merge
        group; rowspan computed by looking at subsequent rows.
      - ``vMerge`` with no val → continuation; dropped from output and
        skipped during logical-column tracking.
    """
    rows = table._element.findall(_QN_TR)

    # Pre-parse each tc -> (vmerge_state, colspan, text)
    def _parse_tc(tc):
        tc_pr = tc.find(_QN_TC_PR)
        colspan = 1
        vmerge = None
        if tc_pr is not None:
            gs = tc_pr.find(_QN_GRID_SPAN)
            if gs is not None:
                val = gs.get(qn("w:val"))
                try:
                    colspan = max(1, int(val)) if val else 1
                except (TypeError, ValueError):
                    colspan = 1
            vm = tc_pr.find(_QN_V_MERGE)
            if vm is not None:
                vmerge = "restart" if vm.get(qn("w:val")) == "restart" else "continue"
        text = "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
        return vmerge, colspan, text

    parsed_rows: list[list[tuple[object, str | None, int, str]]] = []
    for tr in rows:
        tcs = tr.findall(_QN_TC)
        row_parsed: list[tuple[object, str | None, int, str]] = []
        for tc in tcs:
            vmerge, colspan, text = _parse_tc(tc)
            row_parsed.append((tc, vmerge, colspan, text))
        parsed_rows.append(row_parsed)

    # Build a logical-column map: each tc occupies `colspan` consecutive
    # columns; covered columns get None.
    logical: list[list[tuple[object, str | None, int, str] | None]] = []
    for row in parsed_rows:
        lr: list[tuple[object, str | None, int, str] | None] = []
        for entry in row:
            lr.append(entry)
            _, _, colspan, _ = entry
            for _ in range(colspan - 1):
                lr.append(None)
        logical.append(lr)

    n_rows = len(parsed_rows)
    out: list[list[CellInfo]] = []
    for r in range(n_rows):
        out_row: list[CellInfo] = []
        # Iterate logical columns; emit one CellInfo per "owning" tc
        seen_in_row: set[int] = set()
        for col_idx, entry in enumerate(logical[r]):
            if entry is None:
                continue
            tc, vmerge, colspan, text = entry
            if id(tc) in seen_in_row:
                continue
            seen_in_row.add(id(tc))

            if vmerge == "continue":
                continue  # drop continuation cells

            rowspan = 1
            if vmerge == "restart":
                for r2 in range(r + 1, n_rows):
                    if col_idx >= len(logical[r2]):
                        break
                    below = logical[r2][col_idx]
                    if below is None:
                        break
                    _, below_vm, _, _ = below
                    if below_vm == "continue":
                        rowspan += 1
                    else:
                        break

            out_row.append(CellInfo(text=text, colspan=colspan, rowspan=rowspan))
        out.append(out_row)
    return out


def _extract_images_from_paragraph(
    paragraph, doc, out_dir: Path, counter: int
) -> tuple[list[str], list[AssetRef], int]:
    """Walk the paragraph's runs for ``<a:blip>`` embeds; save each.

    Returns ``(markdown_lines, asset_refs, new_counter)``. Empty lists if
    no images. Every markdown ``![](assets/...)`` link is paired with an
    ``AssetRef`` so the deterministic core's H5 check (markdown refs must
    be registered assets) is satisfied.
    """
    md: list[str] = []
    assets: list[AssetRef] = []
    related = doc.part.related_parts
    for blip in paragraph._element.iter(_QN_BLIP):
        rid = blip.get(_QN_EMBED)
        if rid is None or rid not in related:
            continue
        try:
            blob = related[rid].blob
        except Exception:
            continue
        counter += 1
        rel = save_image(blob, out_dir, prefix="img", index=counter)
        if rel is not None:
            md.append(f"![]({rel})")
            assets.append(AssetRef(
                kind="image", rel_path=rel,
                page=1, sha256=sha256_bytes(blob), alt="",
            ))
    return md, assets, counter


def _iter_blocks(doc):
    """Yield Paragraph / Table objects in document order."""
    body = doc.element.body
    para_index = {p._element: p for p in doc.paragraphs}
    tbl_index = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == _QN_P and child in para_index:
            yield para_index[child]
        elif child.tag == _QN_TBL and child in tbl_index:
            yield tbl_index[child]


def _set_language(node: SectionNode, lang: str) -> SectionNode:
    return SectionNode(
        node_id=node.node_id, title=node.title, level=node.level,
        page_start=node.page_start, page_end=node.page_end,
        anchor=node.anchor, language=lang,
        children=tuple(_set_language(c, lang) for c in node.children),
    )


class DocxV2Adapter:
    name = "docx_v2"
    version = "0.8.0"
    extensions = (".docx",)

    def extract(self, src: Path, out_dir_tmp: Path) -> ExtractionResult:
        doc = _docx.Document(str(src))

        md_lines: list[str] = []
        sha = hashlib.sha256(src.read_bytes()).hexdigest()
        md_lines.append(
            f"<!-- generated by kb-extract; do not edit; source_sha256: {sha} -->"
        )
        md_lines.append("")

        children: list[SectionNode] = []
        tables: list[TableRef] = []
        assets: list[AssetRef] = []
        warnings: list[str] = []
        anchor_counter = 0
        table_counter = 0
        image_counter = 0
        body_text_acc: list[str] = []

        def next_anchor() -> str:
            nonlocal anchor_counter
            anchor_counter += 1
            return f"sec-{anchor_counter:04d}"

        section_stack: list[SectionNode] = []
        node_counter = 0

        for block in _iter_blocks(doc):
            if block.__class__.__name__ == "Paragraph":
                style = block.style.name if block.style else ""
                level = _HEADING_LEVELS.get(style)
                if level:
                    node_counter += 1
                    anchor = next_anchor()
                    new_section = SectionNode(
                        node_id=f"{node_counter:04d}",
                        title=block.text.strip() or "(untitled)",
                        level=level,
                        page_start=1,
                        page_end=1,
                        anchor=anchor,
                        language="und",
                        children=(),
                    )
                    md_lines.append(f'<a id="{anchor}"></a>')
                    md_lines.append(f"{'#' * level} {block.text.strip()}")
                    md_lines.append("")

                    while section_stack and section_stack[-1].level >= level:
                        section_stack.pop()

                    if section_stack:
                        parent = section_stack[-1]
                        section_stack[-1] = SectionNode(
                            node_id=parent.node_id, title=parent.title, level=parent.level,
                            page_start=parent.page_start, page_end=parent.page_end,
                            anchor=parent.anchor, language=parent.language,
                            children=tuple([*parent.children, new_section]),
                        )
                    else:
                        children.append(new_section)

                    section_stack.append(new_section)
                else:
                    if style and style not in {"Normal", "Default Paragraph Font"} \
                            and style not in _HEADING_LEVELS:
                        warnings.append(f"docx.unknown_style:{style}")
                    text = block.text.strip()
                    if text:
                        md_lines.append(block.text)
                        md_lines.append("")
                        body_text_acc.append(block.text)
                    # Extract any inline images in this paragraph (after text)
                    img_lines, img_assets, image_counter = _extract_images_from_paragraph(
                        block, doc, out_dir_tmp, image_counter,
                    )
                    assets.extend(img_assets)
                    for img_md in img_lines:
                        md_lines.append(img_md)
                        md_lines.append("")
            elif block.__class__.__name__ == "Table":
                table_counter += 1
                t_anchor = f"tbl-{table_counter:04d}"
                grid = _build_cell_grid(block)
                # Build rows_json (one tuple per logical row, phantom-free,
                # so downstream consumers still see structure even if they
                # can't render rowspan/colspan).
                rows_json: tuple[tuple[str, ...], ...] = tuple(
                    tuple(c.text for c in row) for row in grid
                )
                md_lines.append(f'<a id="{t_anchor}"></a>')
                md_lines.append(cells_to_html(grid))
                md_lines.append("")
                tables.append(TableRef(
                    anchor=t_anchor, page=1, rows_json=rows_json,
                    rendered_asset=None,
                ))

        if not children:
            anchor = next_anchor()
            children.append(SectionNode(
                node_id="0001", title=src.stem, level=1,
                page_start=1, page_end=1, anchor=anchor, language="und",
            ))
            md_lines.insert(2, f'<a id="{anchor}"></a>')
            md_lines.insert(3, f"# {src.stem}")
            md_lines.insert(4, "")
            outline_source = "page_fallback"
        else:
            outline_source = "heading_style"

        lang = _detect_lang(" ".join(body_text_acc))
        root = SectionNode(
            node_id="0000", title=src.stem, level=0, page_start=1, page_end=1,
            anchor="", language=lang, children=tuple(children),
        )
        root = _set_language(root, lang)

        markdown = "\n".join(md_lines) + "\n"
        meta = make_meta(
            src=src,
            adapter_name=self.name,
            adapter_version=self.version,
            tool_versions={"python-docx": _DOCX_VERSION},
            outline_source=outline_source,
            warnings=tuple(warnings),
        )
        return ExtractionResult(
            markdown=markdown, index=root, tables=tuple(tables),
            assets=tuple(assets), meta=meta,
        )
